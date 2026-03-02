import os
import io
import re
import base64
import traceback
import json
import requests
import threading
import concurrent.futures
from datetime import datetime, date
from flask import Flask, render_template, request, jsonify, redirect, url_for, flash, abort, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from docx import Document

from google import genai 
from google.genai import types 

app = Flask(__name__)

# =========================================================
# TELA DE RAIO-X (Tratamento de Erros Global)
# =========================================================
@app.errorhandler(Exception)
def handle_exception(e):
    try: 
        db.session.rollback()
    except Exception: 
        pass
    
    from werkzeug.exceptions import HTTPException
    if isinstance(e, HTTPException): 
        return e
        
    return f"""
    <div style="font-family: sans-serif; padding: 20px; background: #262730; color: #E1E4E8; height: 100vh;">
        <h1 style="color: #FF4B4B;">🚨 Falha Crítica Detectada</h1>
        <p>O sistema encontrou um erro interno.</p>
        <div style="background: #1A1C23; padding: 15px; border-radius: 5px; color: #FFC107; font-family: monospace; overflow-x: auto;">
            <b>{type(e).__name__}</b>: {str(e)}
        </div>
        <p style="margin-top: 20px;">Tire um print e envie para análise.</p>
        <a href="/" style="color: #2196F3; text-decoration: none;">⬅️ Tentar Voltar</a>
    </div>
    """, 500

# =========================================================
# CONFIGURAÇÕES DO BANCO DE DADOS E SEGURANÇA
# =========================================================
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'chave-super-secreta-mude-depois')

db_url = os.environ.get('DATABASE_URL', 'sqlite:///clientes.db')
if db_url.startswith("postgres://"): 
    db_url = db_url.replace("postgres://", "postgresql://", 1)

app.config['SQLALCHEMY_DATABASE_URI'] = db_url
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SQLALCHEMY_ENGINE_OPTIONS'] = {
    "pool_pre_ping": True, 
    "pool_recycle": 60, 
    "pool_timeout": 30
}

db = SQLAlchemy(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'
login_manager.login_message = "Por favor, faça login para acessar."

CHAVE_API_GOOGLE = os.environ.get("GEMINI_API_KEY")
client = genai.Client(api_key=CHAVE_API_GOOGLE) if CHAVE_API_GOOGLE else None

CHAVE_OPENROUTER = os.environ.get("OPENAI_API_KEY")

# =========================================================
# MODELOS DO BANCO DE DADOS
# =========================================================
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(50), unique=True, nullable=False)
    password = db.Column(db.String(255), nullable=False)
    role = db.Column(db.String(20), nullable=False, default='admin') 
    expiration_date = db.Column(db.Date, nullable=True)
    alunos = db.relationship('Aluno', backref='responsavel', lazy=True)

class Aluno(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    nome = db.Column(db.String(100), nullable=False)
    curso = db.Column(db.String(100))
    telefone = db.Column(db.String(20))
    ava_login = db.Column(db.String(255), nullable=True) 
    ava_senha = db.Column(db.String(255), nullable=True) 
    data_cadastro = db.Column(db.DateTime, default=datetime.utcnow)
    status = db.Column(db.String(20), default='Produção') 
    valor = db.Column(db.Float, default=70.0) 
    
    documentos = db.relationship('Documento', backref='aluno', lazy=True, cascade="all, delete-orphan")
    temas = db.relationship('TemaTrabalho', backref='aluno', lazy=True, cascade="all, delete-orphan")

class Documento(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    aluno_id = db.Column(db.Integer, db.ForeignKey('aluno.id'), nullable=False)
    nome_arquivo = db.Column(db.String(255), nullable=False)
    dados_arquivo = db.Column(db.LargeBinary, nullable=False) 
    data_upload = db.Column(db.DateTime, default=datetime.utcnow)

class TemaTrabalho(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    aluno_id = db.Column(db.Integer, db.ForeignKey('aluno.id'), nullable=False)
    titulo = db.Column(db.String(255), nullable=True) 
    texto = db.Column(db.Text, nullable=False)
    data_cadastro = db.Column(db.DateTime, default=datetime.utcnow)

class PromptConfig(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    nome = db.Column(db.String(100), nullable=False)
    texto = db.Column(db.Text, nullable=False)
    is_default = db.Column(db.Boolean, default=False)

class RegistroUso(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    data = db.Column(db.DateTime, default=datetime.utcnow)
    modelo_usado = db.Column(db.String(100))

class SiteSettings(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    whatsapp_template = db.Column(
        db.Text, 
        default="Olá {nome}, seu trabalho de {curso} ficou pronto com excelência! 🎉\nO valor acordado foi R$ {valor}.\n\nSegue a minha chave PIX para liberação do arquivo: [SUA CHAVE AQUI]"
    )
    prompt_password = db.Column(db.String(255), nullable=True)
    convert_api_key = db.Column(db.String(255), nullable=True)

class GeracaoTask(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    status = db.Column(db.String(20), default='Pendente') 
    resultado = db.Column(db.Text, nullable=True) 
    modelo_utilizado = db.Column(db.String(100), nullable=True)
    erro = db.Column(db.Text, nullable=True)
    data_criacao = db.Column(db.DateTime, default=datetime.utcnow)

@login_manager.user_loader
def load_user(user_id):
    return db.session.get(User, int(user_id))

# =========================================================
# PROMPT BASE DE FÁBRICA (SEGURANÇA SE BD ESTIVER VAZIO)
# =========================================================
PROMPT_REGRAS_BASE = """VOCÊ AGORA ASSUME A PERSONA DE UM PROFESSOR UNIVERSITÁRIO AVALIADOR EXTREMAMENTE RIGOROSO E DE ALTA EXCELÊNCIA ACADÊMICA."""

# =========================================================
# INICIALIZAÇÃO BLINDADA E MIGRAÇÕES
# =========================================================
with app.app_context():
    db.create_all()
    
    try: 
        db.session.execute(db.text("ALTER TABLE aluno ADD COLUMN status VARCHAR(20) DEFAULT 'Produção'"))
        db.session.commit()
    except Exception: 
        db.session.rollback()
        
    try: 
        db.session.execute(db.text("ALTER TABLE aluno ADD COLUMN valor FLOAT DEFAULT 70.0"))
        db.session.commit()
    except Exception: 
        db.session.rollback()

    try: 
        db.session.execute(db.text("ALTER TABLE aluno ADD COLUMN ava_login VARCHAR(255)"))
        db.session.commit()
    except Exception: 
        db.session.rollback()

    try: 
        db.session.execute(db.text("ALTER TABLE aluno ADD COLUMN ava_senha VARCHAR(255)"))
        db.session.commit()
    except Exception: 
        db.session.rollback()

    try: 
        db.session.execute(db.text("ALTER TABLE site_settings ADD COLUMN prompt_password VARCHAR(255)"))
        db.session.commit()
    except Exception: 
        db.session.rollback()
        
    try: 
        db.session.execute(db.text("ALTER TABLE site_settings ADD COLUMN convert_api_key VARCHAR(255)"))
        db.session.commit()
    except Exception: 
        db.session.rollback()

    try:
        if not User.query.filter_by(username='admin').first():
            senha_hash = generate_password_hash('admin123')
            admin_user = User(username='admin', password=senha_hash, role='admin')
            db.session.add(admin_user)
            db.session.commit()
    except Exception: 
        db.session.rollback()
        
    try:
        prompt_padrao = PromptConfig.query.filter_by(is_default=True).first()
        if not prompt_padrao:
            novo_prompt = PromptConfig(nome="Padrão Oficial (Desafio UNIASSELVI)", texto=PROMPT_REGRAS_BASE, is_default=True)
            db.session.add(novo_prompt)
            db.session.commit()
    except Exception: 
        db.session.rollback()

    try:
        if not SiteSettings.query.first():
            db.session.add(SiteSettings())
            db.session.commit()
    except Exception: 
        db.session.rollback()

# =========================================================
# MOTOR NATIVO DE IA E TAREFAS EM BACKGROUND
# =========================================================
def limpar_texto_ia(texto):
    try: 
        texto = re.sub(r'\\u([0-9a-fA-F]{4})', lambda m: chr(int(m.group(1), 16)), texto)
    except Exception: 
        pass
    return texto

def chamar_ia(prompt, nome_modelo):
    is_openrouter = "openrouter/" in nome_modelo.lower() or "/" in nome_modelo
    
    if is_openrouter:
        if not CHAVE_OPENROUTER: 
            raise Exception("A Chave da API do OpenRouter não foi configurada.")
            
        modelo_limpo = nome_modelo.replace("openrouter/", "")
            
        headers = {
            "Authorization": f"Bearer {CHAVE_OPENROUTER}",
            "HTTP-Referer": "https://hubmaster-system.com",
            "X-Title": "HubMaster Premium SaaS",
            "Content-Type": "application/json"
        }
        
        payload = {
            "model": modelo_limpo,
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.7,
            "max_tokens": 8000
        }
        
        try:
            resposta = requests.post(
                "https://openrouter.ai/api/v1/chat/completions", 
                headers=headers, 
                json=payload, 
                timeout=120
            )
            
            if resposta.status_code != 200:
                erro_txt = resposta.text
                try: 
                    erro_txt = resposta.json().get('error', {}).get('message', resposta.text)
                except Exception: 
                    pass
                raise Exception(f"Erro OpenRouter ({resposta.status_code}): {erro_txt}")
                
            dados_ia = resposta.json()
            if 'choices' not in dados_ia or len(dados_ia['choices']) == 0:
                raise Exception("A IA do OpenRouter retornou uma resposta vazia.")
                
            texto_retornado = dados_ia['choices'][0]['message']['content']
            return limpar_texto_ia(texto_retornado)
            
        except requests.exceptions.RequestException as e:
            raise Exception(f"Falha de rede ao contactar o OpenRouter: {str(e)}")
            
    else:
        if not client: 
            raise Exception("A Chave da API nativa do Google não foi configurada.")
            
        try:
            resposta = client.models.generate_content(
                model=nome_modelo, 
                contents=prompt
            )
        except Exception as e:
            raise Exception(f"Erro no Gemini Nativo: {str(e)}")
            
        return limpar_texto_ia(resposta.text)

def executar_geracao_bg(app_instance, task_id, prompt_completo, fila_modelos):
    with app_instance.app_context():
        ultimo_erro = ""
        for modelo in fila_modelos:
            try:
                texto_resposta = chamar_ia(prompt_completo, modelo)
                dicionario = extrair_dicionario(texto_resposta)
                
                tags_preenchidas = sum(1 for v in dicionario.values() if v.strip())
                if tags_preenchidas < 10: 
                    raise Exception(f"A IA {modelo} teve preguiça e gerou poucas tags.")
                
                task_verificar = GeracaoTask.query.get(task_id)
                if task_verificar and task_verificar.status == 'Cancelado':
                    db.session.rollback()
                    return 
                    
                novo_registro = RegistroUso(modelo_usado=modelo)
                db.session.add(novo_registro)
                
                task_verificar.status = 'Concluido'
                task_verificar.resultado = json.dumps(dicionario)
                task_verificar.modelo_utilizado = modelo
                db.session.commit()
                return
                
            except Exception as e:
                ultimo_erro = str(e)
                continue
                
        task_erro = GeracaoTask.query.get(task_id)
        if task_erro and task_erro.status != 'Cancelado':
            task_erro.status = 'Erro'
            task_erro.erro = f"Erro Fatal. Todas as IAs reportaram falha técnica. Último erro: {ultimo_erro}"
            db.session.commit()


# =========================================================
# PROCESSAMENTO DE WORD E TAGS
# =========================================================
def preencher_template_com_tags(arquivo_template, dicionario_dados):
    doc = Document(arquivo_template)
    
    def processar_paragrafo(paragrafo):
        texto_original = paragrafo.text
        tem_tag = False
        
        for marcador, texto_novo in dicionario_dados.items():
            if marcador in texto_original:
                texto_original = texto_original.replace(marcador, str(texto_novo))
                tem_tag = True
                
        if tem_tag:
            titulos_memorial = [
                "Resumo", "Contextualização do desafio", "Análise", 
                "Propostas de solução", "Conclusão reflexiva", 
                "Referências", "Autoavaliação"
            ]
            for t in titulos_memorial:
                if texto_original.strip().startswith(t): 
                    texto_original = texto_original.replace(t, f"**{t}**\n", 1)
                    
            titulos_aspectos = ["Aspecto 1:", "Aspecto 2:", "Aspecto 3:", "Por quê:"]
            for t in titulos_aspectos:
                if t in texto_original: 
                    texto_original = texto_original.replace(t, f"\n**{t}** " if "Por quê:" in t else f"**{t}** ")
                    
            if "?" in texto_original and "Por quê:" not in texto_original:
                partes = texto_original.split("?", 1)
                pergunta = partes[0].strip()
                if 10 < len(pergunta) < 150 and not pergunta.startswith("**"):
                    texto_original = f"**{pergunta}?**\n" + partes[1].lstrip()
                    
            texto_original = texto_original.replace("**\n ", "**\n").replace(":**\n:", ":**\n")
            paragrafo.clear()
            linhas = texto_original.split('\n')
            
            for i, linha in enumerate(linhas):
                partes = linha.split('**')
                for j, parte in enumerate(partes):
                    if parte: 
                        run = paragrafo.add_run(parte)
                        if j % 2 == 1: 
                            run.bold = True
                if i < len(linhas) - 1: 
                    paragrafo.add_run('\n')

    for paragrafo in doc.paragraphs: 
        processar_paragrafo(paragrafo)
        
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs: 
                    processar_paragrafo(paragrafo)

    arquivo_saida = io.BytesIO()
    doc.save(arquivo_saida)
    arquivo_saida.seek(0)
    return arquivo_saida

def extrair_texto_docx(arquivo_bytes):
    doc = Document(io.BytesIO(arquivo_bytes)) if isinstance(arquivo_bytes, bytes) else Document(arquivo_bytes)
    return "\n".join([p.text for p in doc.paragraphs])

def extrair_etapa_5(arquivo_bytes):
    doc = Document(io.BytesIO(arquivo_bytes))
    
    linhas = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    idx_inicio = -1
    
    for i, linha in enumerate(linhas):
        if "Lembre-se também de salvar este documento" in linha:
            idx_inicio = i + 1
            
    if idx_inicio == -1:
        for i in range(len(linhas)-1, -1, -1):
            if "memorial analítico" in linhas[i].lower() and "redação" not in linhas[i].lower():
                idx_inicio = i
                break
                
    if idx_inicio == -1 or idx_inicio >= len(linhas):
        return False, "Não foi possível separar as instruções do texto final. O arquivo pode estar fora do padrão."
        
    linhas_finais = linhas[idx_inicio:]
    
    headers_oficiais = [
        "Resumo", "Contextualização do desafio", "Análise", 
        "Propostas de solução", "Conclusão reflexiva", "Referências", "Autoavaliação"
    ]
    
    blocos = ["Memorial\nAnalítico"]
    
    for linha in linhas_finais:
        linha_limpa = linha.replace('**', '').strip()
        if not linha_limpa: continue
        
        if linha_limpa.lower() == "memorial analítico":
            continue
            
        is_header = False
        for h in headers_oficiais:
            if linha_limpa.startswith(h):
                blocos.append(h)
                resto = linha_limpa[len(h):].strip()
                if resto.startswith('-') or resto.startswith(':'):
                    resto = resto[1:].strip()
                if resto:
                    blocos.append(resto)
                is_header = True
                break
                
        if not is_header:
            blocos.append(linha_limpa)
            
    resultado = "\n\n".join(blocos)
    return True, resultado

def extrair_dicionario(texto_ia):
    chaves = [
        "ASPECTO_1", "POR_QUE_1", "ASPECTO_2", "POR_QUE_2", "ASPECTO_3", "POR_QUE_3", 
        "CONCEITOS_TEORICOS", "ANALISE_CONCEITO_1", "ENTENDIMENTO_TEORICO", "SOLUCOES_TEORICAS", 
        "RESUMO_MEMORIAL", "CONTEXTO_MEMORIAL", "ANALISE_MEMORIAL", "PROPOSTAS_MEMORIAL", 
        "CONCLUSAO_MEMORIAL", "REFERENCIAS_ADICIONAIS", "AUTOAVALIACAO_MEMORIAL"
    ]
    dic = {}
    
    for chave in chaves:
        padrao = rf"\[START_{chave}\](.*?)(?=\[END_{chave}\]|\[START_|$)"
        match = re.search(padrao, texto_ia, re.DOTALL | re.IGNORECASE)
        
        if match:
            trecho = match.group(1).strip()
            while trecho.startswith('**') and trecho.endswith('**') and len(trecho) > 4:
                trecho = trecho[2:-2].strip()
            dic[f"{{{{{chave}}}}}"] = trecho
        else:
            dic[f"{{{{{chave}}}}}"] = "" 
            
    return dic

def gerar_correcao_ia_tags(texto_tema, texto_trabalho, critica, nome_modelo):
    config = PromptConfig.query.filter_by(is_default=True).first()
    regras = config.texto if config else PROMPT_REGRAS_BASE
    
    prompt = f"""Você é um professor avaliador extremamente rigoroso.
    TEMA: {texto_tema}
    TRABALHO ATUAL: {texto_trabalho}
    CRÍTICA RECEBIDA: {critica}
    TAREFA: Reescreva as respostas aplicando as melhorias exigidas na crítica. 
    NUNCA formate a resposta inteira em negrito.
    {regras}"""
    
    try:
        texto_resposta = chamar_ia(prompt, nome_modelo)
        return extrair_dicionario(texto_resposta)
    except Exception as e: 
        raise Exception(f"Falha na IA (Correção): {str(e)}")

# =========================================================
# ROTAS PÚBLICAS (PORTAL DO ALUNO)
# =========================================================
@app.route('/portal', methods=['GET', 'POST'])
def portal():
    aluno = None
    erro = None
    
    if request.method == 'POST':
        telefone_busca = request.form.get('telefone', '')
        tel_limpo = re.sub(r'\D', '', telefone_busca)
        
        todos_alunos = Aluno.query.all()
        for a in todos_alunos:
            if a.telefone and re.sub(r'\D', '', a.telefone) == tel_limpo:
                aluno = a
                break
                
        if not aluno: 
            erro = "Nenhum trabalho encontrado para este número de WhatsApp."
            
    return render_template('portal.html', aluno=aluno, erro=erro)

# =========================================================
# ROTAS DO GERADOR E DE REGERAÇÃO
# =========================================================
MODELOS_DISPONIVEIS = [
    "gemini-2.5-flash",                       
    "google/gemini-2.5-flash",                
    "meta-llama/llama-3.3-70b-instruct",      
    "qwen/qwen-2.5-72b-instruct",             
    "mistralai/mistral-nemo"                  
]

@app.route('/')
@login_required
def index():
    if current_user.role == 'cliente' and current_user.expiration_date and date.today() > current_user.expiration_date: 
        return render_template('expirado.html')
        
    todos_alunos = Aluno.query.filter_by(user_id=current_user.id).order_by(Aluno.id.desc()).all()
    alunos_ativos = [a for a in todos_alunos if a.status != 'Pago']
    prompts = PromptConfig.query.all()
    
    return render_template('index.html', modelos=MODELOS_DISPONIVEIS, alunos=alunos_ativos, prompts=prompts)

@app.route('/api/temas/<int:aluno_id>')
@login_required
def get_temas_aluno(aluno_id):
    temas = TemaTrabalho.query.filter_by(aluno_id=aluno_id).all()
    lista_temas = [{"id": t.id, "titulo": t.titulo, "texto": t.texto} for t in temas]
    return jsonify(lista_temas)

@app.route('/api/extrair_memorial/<int:doc_id>', methods=['GET'])
@login_required
def api_extrair_memorial(doc_id):
    doc = Documento.query.get_or_404(doc_id)
    if doc.aluno.user_id != current_user.id and current_user.role != 'admin':
        return jsonify({"sucesso": False, "erro": "Acesso negado."}), 403
    
    try:
        sucesso, texto_memorial = extrair_etapa_5(doc.dados_arquivo)
        return jsonify({
            "sucesso": sucesso, 
            "texto": texto_memorial if sucesso else "", 
            "erro": texto_memorial if not sucesso else ""
        })
    except Exception as e:
        return jsonify({"sucesso": False, "erro": f"Erro interno ao ler o documento: {str(e)}"})


@app.route('/gerar_rascunho', methods=['POST'])
@login_required
def gerar_rascunho():
    tema = request.form.get('tema')
    modelo_selecionado = request.form.get('modelo')
    prompt_id = request.form.get('prompt_id')
    
    if prompt_id and str(prompt_id).isdigit():
        config = PromptConfig.query.get(int(prompt_id))
    else:
        config = PromptConfig.query.filter_by(is_default=True).first()
        
    texto_prompt = config.texto if config else PROMPT_REGRAS_BASE
    prompt_completo = f"TEMA:\n{tema}\n\n{texto_prompt}"
    
    fila_modelos = [modelo_selecionado] + [m for m in MODELOS_DISPONIVEIS if m != modelo_selecionado]
    
    nova_task = GeracaoTask(user_id=current_user.id, status='Pendente')
    db.session.add(nova_task)
    db.session.commit()
    
    thread = threading.Thread(
        target=executar_geracao_bg, 
        args=(app, nova_task.id, prompt_completo, fila_modelos)
    )
    thread.start()
    
    return jsonify({"sucesso": True, "task_id": nova_task.id})

@app.route('/cancelar_geracao/<int:task_id>', methods=['POST'])
@login_required
def cancelar_geracao(task_id):
    task = GeracaoTask.query.get(task_id)
    if task and task.user_id == current_user.id and task.status == 'Pendente':
        task.status = 'Cancelado'
        db.session.commit()
        return jsonify({"sucesso": True})
    return jsonify({"sucesso": False, "erro": "Tarefa não encontrada ou já concluída."})

@app.route('/status_geracao/<int:task_id>', methods=['GET'])
@login_required
def status_geracao(task_id):
    task = GeracaoTask.query.get(task_id)
    if not task or task.user_id != current_user.id:
        return jsonify({"sucesso": False, "erro": "Tarefa não encontrada."})
        
    if task.status == 'Pendente':
        return jsonify({"status": "Pendente"})
        
    if task.status == 'Erro':
        return jsonify({"status": "Erro", "erro": task.erro})
        
    if task.status == 'Cancelado':
        return jsonify({"status": "Cancelado"})
        
    return jsonify({
        "status": "Concluido", 
        "dicionario": json.loads(task.resultado), 
        "modelo_utilizado": task.modelo_utilizado
    })

@app.route('/regerar_trecho', methods=['POST'])
@login_required
def regerar_trecho():
    try:
        dados = request.json
        if not dados: 
            return jsonify({"sucesso": False, "erro": "Nenhum dado recebido."})

        tema = dados.get('tema', '')
        tag = dados.get('tag', '')
        modelo_selecionado = dados.get('modelo', MODELOS_DISPONIVEIS[0])
        prompt_id = dados.get('prompt_id')
        contexto_atual = dados.get('dicionario', {}) 

        if prompt_id and str(prompt_id).isdigit():
            config = PromptConfig.query.get(int(prompt_id))
        else:
            config = PromptConfig.query.filter_by(is_default=True).first()
            
        texto_prompt = config.texto if config else PROMPT_REGRAS_BASE

        texto_contexto = ""
        for chave, valor in contexto_atual.items():
            if valor and str(valor).strip(): 
                texto_contexto += f"{chave}:\n{valor}\n\n"

        prompt_regeracao = f"""Você é um professor avaliador rigoroso.
TEMA/CASO DO DESAFIO:\n{tema}

CONTEXTO ATUAL DO TRABALHO (Para manter a coerência):\n{texto_contexto}

REGRAS GERAIS E ESTRUTURA:\n{texto_prompt}

TAREFA ESPECÍFICA DE CORREÇÃO:
Reescreva APENAS o trecho da tag {tag}. É OBRIGATÓRIO que faça sentido com o contexto. NÃO inclua as marcações [START_{tag}] ou [END_{tag}]. NUNCA formate a resposta toda em negrito (**). Retorne APENAS o texto limpo."""
        
        fila_modelos = [modelo_selecionado] + [m for m in MODELOS_DISPONIVEIS if m != modelo_selecionado]
        ultimo_erro = ""

        for modelo in fila_modelos:
            try:
                novo_texto = chamar_ia(prompt_regeracao, modelo)
                novo_texto = re.sub(rf"\[/?START_{tag}\]", "", novo_texto, flags=re.IGNORECASE)
                novo_texto = re.sub(rf"\[/?END_{tag}\]", "", novo_texto, flags=re.IGNORECASE).strip()
                
                while novo_texto.startswith('**') and novo_texto.endswith('**') and len(novo_texto) > 4: 
                    novo_texto = novo_texto[2:-2].strip()
                    
                return jsonify({"sucesso": True, "novo_texto": novo_texto, "modelo_utilizado": modelo})
                
            except Exception as e:
                ultimo_erro = str(e)
                continue
                
        return jsonify({"sucesso": False, "erro": f"Todas as IAs falharam ao regerar o trecho. Erro: {ultimo_erro}"})
        
    except Exception as e:
        traceback.print_exc()
        return jsonify({"sucesso": False, "erro": str(e)})

@app.route('/gerar_docx_final', methods=['POST'])
@login_required
def gerar_docx_final():
    try:
        dados = request.json or {}
        aluno_id = dados.get('aluno_id')
        dicionario_editado = dados.get('dicionario', {})
        nome_arquivo = dados.get('nome_arquivo', '').strip()
        
        caminho_padrao = os.path.join(app.root_path, 'TEMPLATE_COM_TAGS.docx')
        with open(caminho_padrao, 'rb') as f: 
            arquivo_memoria = io.BytesIO(f.read())
            
        documento_pronto = preencher_template_com_tags(arquivo_memoria, dicionario_editado)
        arquivo_bytes = documento_pronto.read()
        
        if nome_arquivo:
            if not nome_arquivo.lower().endswith('.docx'):
                nome_arquivo += '.docx'
        else:
            nome_arquivo = f"Trabalho_{datetime.now().strftime('%d%m%Y')}.docx"
            
        if aluno_id:
            novo_doc = Documento(
                aluno_id=aluno_id, 
                nome_arquivo=nome_arquivo, 
                dados_arquivo=arquivo_bytes
            )
            db.session.add(novo_doc)
            
            aluno = Aluno.query.get(aluno_id)
            if aluno and (aluno.status == 'Produção' or aluno.status is None): 
                aluno.status = 'Pendente'
                
            db.session.commit()
            
        return jsonify({
            "sucesso": True, 
            "nome_arquivo": nome_arquivo,
            "arquivo_base64": base64.b64encode(arquivo_bytes).decode('utf-8')
        })
        
    except Exception as e: 
        return jsonify({"sucesso": False, "erro": str(e)})

@app.route('/converter_pdf/<int:doc_id>', methods=['POST'])
@login_required
def converter_pdf(doc_id):
    try:
        doc = Documento.query.get_or_404(doc_id)
        if doc.aluno.user_id != current_user.id and current_user.role != 'admin':
            return jsonify({"sucesso": False, "erro": "Acesso negado."}), 403

        if not doc.nome_arquivo.lower().endswith('.docx'):
            return jsonify({"sucesso": False, "erro": "Apenas arquivos .docx podem ser convertidos."})

        config = SiteSettings.query.first()
        convert_api_key = config.convert_api_key if config else None

        if not convert_api_key:
            return jsonify({
                "sucesso": False, 
                "erro": "Para que as logos da faculdade não fiquem tortas, usamos a ConvertAPI (gratuita). Vá em 'Configurações' no seu painel e cadastre a sua Secret Key."
            })

        response = requests.post(
            f'https://v2.convertapi.com/convert/docx/to/pdf?Secret={convert_api_key}',
            files={'File': (doc.nome_arquivo, doc.dados_arquivo)}
        )
        
        dados_resposta = response.json()
        
        if response.status_code == 200:
            pdf_base64 = dados_resposta['Files'][0]['FileData']
            pdf_bytes = base64.b64decode(pdf_base64)
            
            novo_nome = doc.nome_arquivo.replace('.docx', '.pdf').replace('.DOCX', '.pdf')
            
            novo_doc = Documento(
                aluno_id=doc.aluno_id,
                nome_arquivo=novo_nome,
                dados_arquivo=pdf_bytes
            )
            db.session.add(novo_doc)
            db.session.commit()
            
            return jsonify({"sucesso": True})
        else:
            erro_msg = dados_resposta.get('Message', 'Falha desconhecida.')
            return jsonify({"sucesso": False, "erro": f"A API recusou: {erro_msg}"})
            
    except Exception as e:
        return jsonify({"sucesso": False, "erro": str(e)})

# =========================================================
# GABARITO INTELIGENTE (IA DE ELITE)
# =========================================================
def extrair_json_seguro(texto):
    try:
        match = re.search(r'\[.*\]', texto, re.DOTALL)
        if match:
            return json.loads(match.group(0))
        return json.loads(texto)
    except Exception:
        return []

def consultar_ia_gabarito(texto_prova, modelo):
    prompt = f"""Você é um professor PhD especialista em criar gabaritos perfeitos.
Resolva a prova abaixo. Retorne EXATAMENTE um Array JSON puro, sem formatação markdown (sem ```json), sem texto antes ou depois.
O JSON deve ter esta estrutura exata para cada questão encontrada:
[
  {{"questao": 1, "resposta": "A", "justificativa": "Motivo curto e direto."}},
  {{"questao": 2, "resposta": "C", "justificativa": "Outro motivo."}}
]

PROVA:
{texto_prova}
"""
    try:
        resposta = chamar_ia(prompt, modelo)
        return extrair_json_seguro(resposta)
    except Exception as e:
        print(f"Erro no modelo {modelo}: {e}")
        return []

@app.route('/gabarito_inteligente')
@login_required
def gabarito_inteligente():
    return render_template('gabarito_inteligente.html')

@app.route('/api/gerar_gabarito', methods=['POST'])
@login_required
def api_gerar_gabarito():
    texto_prova = request.form.get('prova', '')
    if not texto_prova:
        return jsonify({"sucesso": False, "erro": "O texto da prova está vazio."})

    # Utilizamos o modelo mais avançado do mundo para esta tarefa
    modelo_elite = "anthropic/claude-3.5-sonnet"

    resultado_ia = consultar_ia_gabarito(texto_prova, modelo_elite)

    if not resultado_ia or len(resultado_ia) == 0:
        return jsonify({"sucesso": False, "erro": "A IA falhou ao processar o gabarito ou a prova é muito extensa. Tente novamente."})

    novo_registro = RegistroUso(modelo_usado=modelo_elite)
    db.session.add(novo_registro)
    db.session.commit()

    gabarito_final = []
    for idx, q in enumerate(resultado_ia):
        letra_crua = str(q.get('resposta', '')).upper().strip()
        match_letra = re.search(r'[A-E]', letra_crua)
        letra = match_letra.group(0) if match_letra else "?"
        
        gabarito_final.append({
            "questao": q.get('questao', idx + 1),
            "resposta": letra,
            "justificativa": q.get('justificativa', 'Justificativa não fornecida.'),
            "confianca": "100%"
        })

    return jsonify({
        "sucesso": True, 
        "gabarito": gabarito_final,
        "modelo_utilizado": modelo_elite
    })

# =========================================================
# REVISÃO AVULSA
# =========================================================
@app.route('/revisao_avulsa')
@login_required
def revisao_avulsa(): 
    return render_template('revisao_avulsa.html', modelos=MODELOS_DISPONIVEIS)

@app.route('/avaliar_avulso', methods=['POST'])
@login_required
def avaliar_avulso():
    tema = request.form.get('tema')
    modelo = request.form.get('modelo')
    arquivo = request.files.get('arquivo_trabalho')
    
    if not arquivo or not arquivo.filename.endswith('.docx'): 
        return jsonify({"erro": "Envie um arquivo .docx válido."}), 400
        
    texto_trabalho = extrair_texto_docx(arquivo.read())
    prompt = f"""Você é um professor avaliador extremamente rigoroso.
Analise o TEMA: {tema} \nE O TRABALHO DO ALUNO: {texto_trabalho}
Faça uma crítica de 3 linhas apontando o que falta para tirar nota máxima. 
REGRA: Seja direto, NUNCA use formatações como negrito (**), itálico, bullet points ou títulos. Responda apenas com texto limpo."""
    
    try: 
        resposta_ia = chamar_ia(prompt, modelo)
        critica_limpa = resposta_ia.replace('*', '').replace('#', '').strip()
        return jsonify({
            "critica": critica_limpa, 
            "texto_extraido": texto_trabalho
        })
    except Exception as e: 
        return jsonify({"erro": str(e)}), 500

@app.route('/corrigir_avulso', methods=['POST'])
@login_required
def corrigir_avulso():
    tema = request.form.get('tema')
    texto_trabalho = request.form.get('texto_extraido')
    critica = request.form.get('critica')
    modelo = request.form.get('modelo')
    
    try:
        respostas = gerar_correcao_ia_tags(tema, texto_trabalho, critica, modelo)
        
        caminho_padrao = os.path.join(app.root_path, 'TEMPLATE_COM_TAGS.docx')
        with open(caminho_padrao, 'rb') as f: 
            arquivo_memoria = io.BytesIO(f.read())
            
        documento_pronto = preencher_template_com_tags(arquivo_memoria, respostas)
        arquivo_base64 = base64.b64encode(documento_pronto.read()).decode('utf-8')
        
        return jsonify({
            "arquivo_base64": arquivo_base64, 
            "nome_arquivo": "Trabalho_Revisado_IA.docx"
        })
    except Exception as e: 
        return jsonify({"erro": str(e)}), 500

# =========================================================
# LOGIN E ADMINISTRAÇÃO
# =========================================================
@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated: 
        return redirect(url_for('index'))
        
    if request.method == 'POST':
        user = User.query.filter_by(username=request.form.get('username')).first()
        if user and check_password_hash(user.password, request.form.get('password')):
            login_user(user)
            return redirect(url_for('index'))
        else: 
            flash('Credenciais incorretas.', 'error')
            
    return render_template('login.html')

@app.route('/logout')
@login_required
def logout(): 
    logout_user()
    return redirect(url_for('login'))

@app.route('/mudar_senha', methods=['GET', 'POST'])
@login_required
def mudar_senha():
    if request.method == 'POST':
        senha_atual = request.form.get('senha_atual')
        nova_senha = request.form.get('nova_senha')
        
        if check_password_hash(current_user.password, senha_atual):
            current_user.password = generate_password_hash(nova_senha)
            db.session.commit()
            flash('Sua senha foi atualizada!', 'success')
            return redirect(url_for('index'))
            
    return render_template('mudar_senha.html')

@app.route('/admin', methods=['GET', 'POST'])
@login_required
def admin():
    if current_user.role not in ['admin', 'sub-admin']: 
        abort(403)
        
    if request.method == 'POST':
        data_expiracao = request.form.get('expiration_date')
        novo_user = User(
            username=request.form.get('username'), 
            password=generate_password_hash(request.form.get('password')),
            role=request.form.get('role'), 
            expiration_date=datetime.strptime(data_expiracao, '%Y-%m-%d').date() if data_expiracao else None
        )
        db.session.add(novo_user)
        db.session.commit()
        flash('Usuário criado com sucesso.', 'success')

    users = User.query.all()
    return render_template('admin.html', users=users, hoje=date.today())

@app.route('/edit_user/<int:id>', methods=['GET', 'POST'])
@login_required
def edit_user(id):
    if current_user.role not in ['admin', 'sub-admin']: 
        abort(403)
        
    user = User.query.get_or_404(id)
    if request.method == 'POST':
        nova_senha = request.form.get('password')
        if nova_senha: 
            user.password = generate_password_hash(nova_senha)
            
        if current_user.role == 'admin': 
            user.role = request.form.get('role')
            
        data_expiracao = request.form.get('expiration_date')
        user.expiration_date = datetime.strptime(data_expiracao, '%Y-%m-%d').date() if data_expiracao else None
        
        db.session.commit()
        flash('Usuário atualizado com sucesso!', 'success')
        return redirect(url_for('admin'))
        
    return render_template('edit_user.html', user=user)

@app.route('/delete_user/<int:id>')
@login_required
def delete_user(id):
    if current_user.role not in ['admin', 'sub-admin']: 
        abort(403)
        
    user = User.query.get_or_404(id)
    db.session.delete(user)
    db.session.commit()
    return redirect(url_for('admin'))


if __name__ == '__main__':
    app.run(debug=True)
