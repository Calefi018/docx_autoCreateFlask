import io
import os
import re
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

def preencher_template_com_tags(arquivo_template, dicionario_dados):
    """
    MOTOR 1: Usado APENAS para os Trabalhos Acadêmicos gerados por IA.
    Aplica formatação Markdown (**negrito**) e regras específicas de títulos.
    """
    doc = Document(arquivo_template)
    
    def processar_paragrafo(paragrafo):
        texto_original = paragrafo.text
        tem_tag = False
        
        for marcador, texto_novo in dicionario_dados.items():
            if marcador in texto_original:
                texto_original = texto_original.replace(marcador, str(texto_novo))
                tem_tag = True
                
        if tem_tag:
            titulos_memorial = [
                "Resumo", "Contextualização do desafio", "Análise", 
                "Propostas de solução", "Conclusão reflexiva", 
                "Referências", "Autoavaliação"
            ]
            for t in titulos_memorial:
                if texto_original.strip().startswith(t): 
                    texto_original = texto_original.replace(t, f"**{t}**\n", 1)
                    
            titulos_aspectos = ["Aspecto 1:", "Aspecto 2:", "Aspecto 3:", "Por quê:"]
            for t in titulos_aspectos:
                if t in texto_original: 
                    texto_original = texto_original.replace(t, f"\n**{t}** " if "Por quê:" in t else f"**{t}** ")
                    
            if "?" in texto_original and "Por quê:" not in texto_original:
                partes = texto_original.split("?", 1)
                pergunta = partes[0].strip()
                if 10 < len(pergunta) < 150 and not pergunta.startswith("**"):
                    texto_original = f"**{pergunta}?**\n" + partes[1].lstrip()
                    
            texto_original = texto_original.replace("**\n ", "**\n").replace(":**\n:", ":**\n")
            paragrafo.clear()
            linhas = texto_original.split('\n')
            
            for i, linha in enumerate(linhas):
                partes = linha.split('**')
                for j, parte in enumerate(partes):
                    if parte: 
                        run = paragrafo.add_run(parte)
                        if j % 2 == 1: 
                            run.bold = True
                if i < len(linhas) - 1: 
                    paragrafo.add_run('\n')

    for paragrafo in doc.paragraphs: 
        processar_paragrafo(paragrafo)
        
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs: 
                    processar_paragrafo(paragrafo)

    arquivo_saida = io.BytesIO()
    doc.save(arquivo_saida)
    arquivo_saida.seek(0)
    return arquivo_saida


def preencher_template_extensao(arquivo_template, dicionario_dados):
    """
    MOTOR 2 (ALTA PRECISÃO): Usado APENAS para Projetos de Extensão.
    Filtra Caixas de Texto, junta tags fragmentadas pelo Word e auto-corrige erros.
    """
    doc = Document(arquivo_template)

    def arrumar_tag_quebrada(match):
        # Remove lixos invisíveis e espaços (ex: {{ DATA_ 08 }} vira {{DATA_8}})
        inner = match.group(1).replace(' ', '').replace('\u200b', '').replace('\xa0', '').upper()
        inner = inner.replace('-', '_')
        
        if inner == 'NOMEALUNO': return '{{NOME_ALUNO}}'
        if inner == 'MATRICULA': return '{{MATRICULA}}'
        if inner == 'CURSO': return '{{CURSO}}'
        
        # Converte zeros extras (DATA_08 -> DATA_8)
        if inner.startswith('DATA'):
            num = inner.replace('DATA', '').replace('_', '')
            if num.isdigit():
                return f"{{{{DATA_{int(num)}}}}}"
                
        return f"{{{{{inner}}}}}"

    def has_shape(run):
        # Verifica se este 'run' é, na verdade, uma imagem ou o "esqueleto" da Caixa de Texto
        for tag in ['w:drawing', 'w:pict', 'w:object', 'v:shape']:
            try:
                if len(list(run._element.iter(qn(tag)))) > 0:
                    return True
            except:
                pass
        return False

    def substituir_em_paragrafo(p):
        # 1. Isola apenas os pedaços de texto puro, protegendo os esqueletos das caixas de texto
        safe_runs = [r for r in p.runs if not has_shape(r)]
        if not safe_runs:
            return
            
        # 2. Junta os textos que o Word quebrou invisivelmente
        full_text = "".join(r.text for r in safe_runs if r.text)
        if "{{" not in full_text:
            return

        # 3. Passa a vassoura nos espaços vazios e erros de digitação das chaves
        texto_limpo = re.sub(r'[\u200b\u200c\u200d\ufeff\xa0]', '', full_text)
        texto_limpo = re.sub(r'\{\{(.*?)\}\}', arrumar_tag_quebrada, texto_limpo)
        
        modificou = False
        for marcador, valor in dicionario_dados.items():
            if marcador in texto_limpo:
                texto_limpo = texto_limpo.replace(marcador, str(valor))
                modificou = True
                
        # 4. Injeta a data limpa no primeiro espaço de texto e apaga os pedaços quebrados
        if modificou or texto_limpo != full_text:
            safe_runs[0].text = texto_limpo
            for r in safe_runs[1:]:
                r.text = ""

    # RAIO-X: Entra nas Caixas de Texto, Tabelas Flutuantes e Textos Normais
    for node in doc.element.body.iter():
        if isinstance(node, CT_P):
            p = Paragraph(node, doc)
            substituir_em_paragrafo(p)

    # RAIO-X: Entra nos Cabeçalhos e Rodapés
    for section in doc.sections:
        if section.header:
            for node in section.header._element.iter():
                if isinstance(node, CT_P):
                    substituir_em_paragrafo(Paragraph(node, section.header))
        if section.footer:
            for node in section.footer._element.iter():
                if isinstance(node, CT_P):
                    substituir_em_paragrafo(Paragraph(node, section.footer))

    arquivo_saida = io.BytesIO()
    doc.save(arquivo_saida)
    arquivo_saida.seek(0)
    return arquivo_saida


def extrair_texto_docx(arquivo_bytes):
    doc = Document(io.BytesIO(arquivo_bytes)) if isinstance(arquivo_bytes, bytes) else Document(arquivo_bytes)
    return "\n".join([p.text for p in doc.paragraphs])


def extrair_etapa_5(arquivo_bytes):
    doc = Document(io.BytesIO(arquivo_bytes))
    linhas = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    idx_inicio = -1
    
    for i, linha in enumerate(linhas):
        if "Lembre-se também de salvar este documento" in linha:
            idx_inicio = i + 1
            
    if idx_inicio == -1:
        for i in range(len(linhas)-1, -1, -1):
            if "memorial analítico" in linhas[i].lower() and "redação" not in linhas[i].lower():
                idx_inicio = i
                break
                
    if idx_inicio == -1 or idx_inicio >= len(linhas):
        return False, "Não foi possível separar as instruções do texto final. O arquivo pode estar fora do padrão."
        
    linhas_finais = linhas[idx_inicio:]
    headers_oficiais = [
        "Resumo", "Contextualização do desafio", "Análise", 
        "Propostas de solução", "Conclusão reflexiva", "Referências", "Autoavaliação"
    ]
    
    blocos = ["Memorial\nAnalítico"]
    
    for linha in linhas_finais:
        linha_limpa = linha.replace('**', '').strip()
        if not linha_limpa: continue
        
        if linha_limpa.lower() == "memorial analítico":
            continue
            
        is_header = False
        for h in headers_oficiais:
            if linha_limpa.startswith(h):
                blocos.append(h)
                resto = linha_limpa[len(h):].strip()
                if resto.startswith('-') or resto.startswith(':'):
                    resto = resto[1:].strip()
                if resto:
                    blocos.append(resto)
                is_header = True
                break
                
        if not is_header:
            blocos.append(linha_limpa)
            
    resultado = "\n\n".join(blocos)
    return True, resultado
